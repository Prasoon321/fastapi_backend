# import os
# import hashlib
import docx
# from docx import Document
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from fastapi import FastAPI, UploadFile, Form, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.llms import HuggingFaceHub
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain_pinecone import PineconeVectorStore
load_dotenv()
app = FastAPI()
origins = [
    "https://www.smartdocsai.site",
      "http://localhost:5173",
      "http://127.0.0.1:5173",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
index_name = "docquery"  # Pinecone index name

def extract_text_from_word(word_file: UploadFile):
    # print("entered the word ")

    try:
        # Load the Word document
        doc = docx.Document(word_file.file)
        text = ""
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                text += f"{paragraph_text} "
        
        # Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
        final_documents = text_splitter.split_text(text)
        documents = [Document(page_content=chunk) for chunk in final_documents]
        # print(documents)
        return documents
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from Word document: {str(e)}")

# Extract text from plain text (.txt) files
def extract_text_from_txt(txt_file: UploadFile):
    # print("entered the txt ")

    try:
        # Read the plain text file
        text = txt_file.file.read().decode('utf-8').strip()

        # Split the text into manageable chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
        final_documents = text_splitter.split_text(text)
        documents = [Document(page_content=chunk) for chunk in final_documents]
        # print(documents)
        return documents
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from TXT file: {str(e)}")



def extract_text_from_pdf(pdf_file: UploadFile):
    # print("entered pdf text extracter")
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file.file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            text += page.extract_text()
            if page_text:
                page_text = page_text.replace("\n", " ").strip()
                text += page_text
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
        final_documents = text_splitter.split_text(text)
        documents = [Document(page_content=chunk) for chunk in final_documents]
        return documents
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

@app.post("/api/upload-file")
async def upload_file(file: UploadFile = File(...)):
    try:
        # print("entered the api ")
        # Check if file is provided
        if not file:
            raise HTTPException(status_code=400, detail="Missing file")

        # Extract the file extension
        file_extension = file.filename.split(".")[-1].lower()

        # Route the file to appropriate text extraction function
        if file_extension == "pdf":
            documents = extract_text_from_pdf(file)
        elif file_extension in ["doc", "docx"]:
            documents = extract_text_from_word(file)
        elif file_extension == "txt":
            documents = extract_text_from_txt(file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF, Word document, or plain text file.")

        # Embedding and storing the documents (kept from your original code)
                    # BAAI/bge-base-en-v1.5
            # sentence-transformers/all-MiniLM-L12-v2 
            # sentence-transformers/all-MiniLM-L6-v2 
            # sentence-transformers/all-distilroberta-v1
            # sentence-transformers/paraphrase-MiniLM-L12-v2
        huggingface_embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L12-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        docsearch = PineconeVectorStore.from_documents(
            documents=documents,
            embedding=huggingface_embeddings,
            index_name=index_name
        )

        return JSONResponse(content={"message": "File uploaded and content stored successfully"})
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error uploading file: {str(e)}")
@app.post("/api/query-pinecone")
async def query_pinecone(query: str = Form(...)):
    try:
        if not query:
            raise HTTPException(status_code=400, detail="Missing query")
        huggingface_embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L12-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        docsearch = PineconeVectorStore(
            embedding=huggingface_embeddings,
            index_name=index_name
        )
        relevant_documents = docsearch.similarity_search(query, k=1)
        doc_data = [doc.page_content for doc in relevant_documents]
        return JSONResponse(content={"answer": doc_data})
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error querying Pinecone: {str(e)}")
@app.get("/")
async def root():
    return {"message": "Hello World"}


 